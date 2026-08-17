from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path
from datetime import date

OUT = Path(r"C:\Users\user\OneDrive - 상지대학교\문서\ChatGPT\마케팅\manuals\브링케어_네이버블로그_마스터매뉴얼_v1.1.docx")
LOGO = Path(r"C:\Users\user\OneDrive - 상지대학교\문서\ChatGPT\회사 체계\FM\signage-template\bring-care-logo.png")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "0A2A66"
BLUE = "2368B5"
SKY = "DCEBFA"
MINT = "E2F2EB"
BEIGE = "F7F0E3"
YELLOW = "FFF3BF"
LIGHT = "F3F6FA"
GRAY = "606B7A"
RED = "B3261E"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.72)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.32)
sec.footer_distance = Inches(0.32)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_font(run, size=10.2, bold=False, color="20252B", name="Malgun Gothic"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def set_p(p, before=0, after=5, line=1.28, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Malgun Gothic"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.font.size = Pt(10.2)
normal.font.color.rgb = RGBColor.from_string("20252B")
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.28

for name, size, color, before, after in [
    ("Title", 30, NAVY, 0, 10),
    ("Subtitle", 13, GRAY, 0, 8),
    ("Heading 1", 20, NAVY, 16, 8),
    ("Heading 2", 14.5, BLUE, 12, 5),
    ("Heading 3", 11.5, NAVY, 9, 3),
]:
    st = styles[name]
    st.font.name = "Malgun Gothic"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    st.font.size = Pt(size)
    st.font.bold = name != "Subtitle"
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Manual Label" not in styles:
    st = styles.add_style("Manual Label", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "Malgun Gothic"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    st.font.size = Pt(8.5)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLUE)
    st.paragraph_format.space_after = Pt(3)

def add_header_footer(section):
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("BRING Care  |  NAVER BLOG OPERATIONS MANUAL")
    set_font(r, 8, True, GRAY)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("브링케어 내부 운영 문서  ·  최신 규칙은 개정 이력 확인  ·  ")
    set_font(r, 8, False, GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)

add_header_footer(sec)

def add_rule(color=SKY, size=8):
    p = doc.add_paragraph()
    set_p(p, 2, 8)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

def add_para(text, bold_prefix=None, after=5, color="20252B", align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_p(p, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, 10.2, True, NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, 10.2, False, color)
    else:
        r = p.add_run(text)
        set_font(r, 10.2, False, color)
        r.italic = italic
    return p

def add_bullets(items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        set_p(p, after=3)
        r = p.add_run(item)
        set_font(r, 10, False)

def new_numbering_id():
    numbering = doc.part.numbering_part.element
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), "7")
    num.append(abs_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id

def add_numbers(items):
    num_id = new_numbering_id()
    for item in items:
        p = doc.add_paragraph()
        ppr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        ppr.append(num_pr)
        set_p(p, after=3)
        r = p.add_run(item)
        set_font(r, 10, False)

def add_callout(title, body, fill=MINT, accent=BLUE):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.7)
    cell = t.cell(0, 0)
    cell.width = Inches(6.7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 180, 210, 180, 210)
    p = cell.paragraphs[0]
    set_p(p, after=4)
    r = p.add_run(title)
    set_font(r, 10.5, True, accent)
    p2 = cell.add_paragraph()
    set_p(p2, after=0)
    r2 = p2.add_run(body)
    set_font(r2, 9.8, False)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

def add_table(headers, rows, widths=None, font_size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_shading(c, NAVY)
        set_cell_margins(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p(p, after=0)
        r = p.add_run(h)
        set_font(r, font_size, True, WHITE)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            set_cell_margins(c)
            if ri % 2 == 1:
                set_cell_shading(c, LIGHT)
            p = c.paragraphs[0]
            set_p(p, after=0, line=1.15)
            r = p.add_run(str(val))
            set_font(r, font_size, False)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def chapter(num, title, purpose=None):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    set_p(p, before=0, after=3)
    r = p.add_run(f"CHAPTER {num:02d}")
    set_font(r, 9, True, BLUE)
    h = doc.add_heading(title, level=1)
    if purpose:
        add_callout("이 장의 목적", purpose, fill=SKY)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

# Cover
for _ in range(3):
    doc.add_paragraph()
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(1.45))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_p(p, before=18, after=8)
r = p.add_run("브링케어 네이버 블로그")
set_font(r, 14, True, BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_p(p, after=10)
r = p.add_run("마스터 운영 매뉴얼")
set_font(r, 30, True, NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_p(p, after=28)
r = p.add_run("시장조사 · 콘텐츠 기획 · 원고 작성 · 사진 · 편집 · 검증 · 발행 · 성과개선")
set_font(r, 12.5, False, GRAY)
add_rule(BLUE, 14)
for _ in range(3):
    doc.add_paragraph()
add_callout("문서의 용도", "브링케어 블로그를 누가 작성하더라도 같은 판단과 품질로 운영하도록 만든 살아 있는 기준서입니다. 새로운 합의가 생기면 기존 규칙을 지우지 말고 개정 이력과 해당 장에 반영합니다.", fill=BEIGE, accent=NAVY)
add_para("버전 1.0  |  기준일 2026-08-17  |  문서 소유자 BRING Care", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
doc.add_page_break()

doc.add_heading("문서 관리 원칙", level=1)
add_table(["항목", "운영 기준"], [
    ("단일 기준 문서", "규칙은 이 파일에 누적한다. 메신저 합의만 남기지 않는다."),
    ("개정 방식", "새 규칙을 해당 장에 반영하고 맨 뒤 개정 이력에 날짜·사유·영향 범위를 기록한다."),
    ("우선순위", "최신 개정 내용 > 본 매뉴얼 본문 > 과거 초안·대화 기록 순으로 적용한다."),
    ("사실과 취향", "법·정책·제품·안전·가격은 공식 근거로 검증하고, 편집 취향은 브랜드 규칙으로 관리한다."),
    ("예외", "예외가 반복되면 임시 대응으로 두지 말고 정식 규칙으로 승격한다."),
], widths=[1.5, 5.2], font_size=9)
h2("권장 파일명 규칙")
add_bullets([
    "마스터 파일: 브링케어_네이버블로그_마스터매뉴얼_v1.1.docx",
    "대규모 체계 변경: v2.0처럼 주 버전을 올린다.",
    "문구·체크리스트 보완: v1.1, v1.2처럼 부 버전을 올린다.",
    "발행 원고: blog/YYYY-MM-DD-slug.md",
    "사진 자산: blog/assets/YYYY-MM-DD-slug/",
    "자동화 상태: blog/automation/backlog.md",
])
h2("목차")
toc = [
"브랜드와 사업의 정확한 정의", "블로그 전략과 콘텐츠 포트폴리오", "독자·검색의도·콘텐츠 역할", "시장조사와 핫키워드 발굴", "주제 승인 게이트", "글 유형별 설계", "제목과 대표 이미지", "도입부 설계", "본문 설득 구조", "네이버 모바일 편집 규칙", "사진 수집·저작권·개인정보", "공식 자료 활용과 출처 처리", "브링케어 연결 문단과 CTA", "현장 작업 글 작성법", "검색 해결형 글 작성법", "대중 유입형·홈피드형 글 작성법", "부동산·연예인 공간·시사 인접 주제", "쿠팡 파트너스 수익형", "금지 표현·안전·사실성", "발행 전 검증", "네이버 편집기 발행 절차", "발행 후 공개 검수", "성과 측정과 개선", "자동화 운영", "실패 사례와 교정법", "표준 입력 양식", "표준 출력 패키지", "체크리스트·기록지·개정 이력"
]
add_numbers([f"{i+1:02d}. {x}" for i, x in enumerate(toc)])

chapter(1, "브랜드와 사업의 정확한 정의", "글을 많이 쓰는 것보다 먼저, 브링케어가 무엇을 하고 무엇을 하지 않는지 모든 작성자가 동일하게 이해한다.")
h2("1.1 한 문장 포지셔닝")
add_callout("브링케어란", "원주 원룸·다가구·소형 상가에서 반복되는 관리 업무를 한 창구로 접수하고, 현장 확인·분류·건물주 승인·일정 조율·완료 확인·사진 보고까지 연결하는 건물관리 서비스입니다.", fill=MINT)
h2("1.2 역할 경계")
add_table(["구분", "브링케어가 하는 일", "주의할 표현"], [
    ("직접 관리", "문의 접수, 현장 확인, 문제 분류, 사진 기록, 건물주 보고, 일정 조율, 완료 확인", "직접 시공하지 않은 일을 ‘수리했습니다’라고 쓰지 않는다."),
    ("외부 전문업체", "전기·가스·구조·제조사 AS·전문 보수의 실제 작업", "브링케어가 시공 주체인 것처럼 오인시키지 않는다."),
    ("건물주 결정", "비용 승인, 작업 범위, 보수 우선순위, 계약 범위", "승인 전 일을 확정된 것처럼 표현하지 않는다."),
    ("미확인", "누가 수행했는지, 원인이 무엇인지, 완료 여부가 불명확한 항목", "추측하지 않고 ‘확인 필요’로 남긴다."),
], widths=[1.1, 3.3, 2.3], font_size=8.5)
h2("1.3 고정 정보")
add_bullets(["전화 010-6566-3603", "카카오채널 BRING Care", "블로그 https://blog.naver.com/bringcare", "YouTube @bring_engineering", "Instagram @bring._.care"])
h2("1.4 글에서 보여줘야 하는 차별성")
add_bullets([
    "전화번호 전달이 아니라 확인 → 분류 → 승인 → 조율 → 완료 확인까지 이어진다.",
    "무조건 공사로 연결하지 않고 안전성·생활 영향·시급성·비용을 보고 순서를 정한다.",
    "건물주가 현장에 매번 가지 않아도 사진과 처리 상태를 확인할 수 있게 한다.",
    "청소·공실·임차인 문의·시설 문제를 각각 따로 관리하는 수고를 한 창구로 줄인다.",
])

chapter(2, "블로그 전략과 콘텐츠 포트폴리오", "조회수만 노리는 글과 서비스 소개만 하는 글을 분리하지 않고, 유입에서 상담까지 이어지는 포트폴리오로 운영한다.")
h2("2.1 네 가지 콘텐츠 역할")
add_table(["역할", "독자 상태", "글의 임무", "대표 근거"], [
    ("유입", "처음 방문", "생활·주거 문제의 빠르고 유용한 답 제공", "공식자료, 실용 체크리스트"),
    ("검토", "업체 비교 중", "관리 부담과 반대 의견을 풀고 판단 기준 제공", "비교표, 업무 흐름"),
    ("증거", "실제 수행 여부 확인", "현장 과정과 결과를 사진으로 입증", "실제 사진, 작업 기록"),
    ("신뢰", "계약 전 위험 검토", "역할 경계·한계·원칙을 투명하게 설명", "정책, 업무 범위"),
], widths=[0.8, 1.25, 3.0, 1.7], font_size=8.4)
h2("2.2 월간 기본 비율")
add_callout("포트폴리오 기본값", "유입 60% · 현장증명 30% · 전환 10%. 이 비율은 한 글의 문단 비율이 아니라 한 달 전체 게시물 구성입니다.", fill=YELLOW, accent=NAVY)
h2("2.3 한 글의 단일성")
add_bullets(["한 독자", "한 문제", "한 약속", "한 주역할", "본문의 주 CTA 하나"])
add_para("모든 것을 한 글에 넣으면 누구에게도 선명하지 않습니다. 회사소개, 여러 서비스, 전화·카카오·저장·관련글 행동을 한꺼번에 요구하지 않습니다.")

chapter(3, "독자·검색의도·콘텐츠 역할", "‘누가, 언제, 무엇 때문에 검색했는가’를 먼저 정하고 글의 답과 행동을 일치시킨다.")
h2("3.1 핵심 독자군")
add_table(["독자", "대표 장면", "불안", "원하는 답"], [
    ("자가 거주자", "냄새·오염·가전 문제 발견", "직접 해도 되는지", "안전한 확인 순서"),
    ("원룸 임차인", "시설 민원 발생", "누구에게 말해야 하는지", "증상 기록·연락 순서"),
    ("원룸·다가구 건물주", "공실·민원·청소 반복", "시간·확인 부담", "누가 확인·조율하는지"),
    ("타지역 건물주", "현장 방문 어려움", "처리 상태 불투명", "사진 보고와 완료 확인"),
    ("소형 상가 운영자", "공용부·시설 문제", "영업 영향", "우선순위와 일정 조율"),
], widths=[1.15, 2.0, 1.7, 1.85], font_size=8.3)
h2("3.2 검색 의도 네 가지")
add_bullets([
    "정보형: ‘방법’, ‘원인’, ‘확인 순서’를 찾는다.",
    "비교형: 직접 처리와 업체 의뢰, 청소와 수리의 차이를 찾는다.",
    "행동형: 신청·예약·문의·준비물처럼 바로 실행할 방법을 찾는다.",
    "지역 서비스형: 원주 건물관리·청소·시설관리 업체를 찾는다.",
])
h2("3.3 브리프 한 줄 공식")
add_callout("독자 브리프", "[독자]가 [상황]에서 느끼는 [불안]을 풀고, 글을 다 읽으면 [약속한 답]을 얻어 [하나의 행동]을 하게 한다.", fill=BEIGE)

chapter(4, "시장조사와 핫키워드 발굴", "유행어를 억지로 붙이지 않고 실제 관심·검색·행동 의도와 브링케어 연결성이 동시에 있는 빈틈을 찾는다.")
h2("4.1 조사 순서")
add_numbers([
    "네이버 크리에이터 어드바이저에서 최근 유입 콘텐츠와 연령·성별 반응을 확인한다.",
    "네이버 데이터랩에서 계절성·상승 흐름·연관 검색어를 확인한다.",
    "블랙키위·판다랭크·리뷰언즈·데이터랩툴즈 등은 후보 탐색과 경쟁 강도 참고에 사용한다.",
    "네이버 검색 결과에서 상위 글 10개를 직접 열어 제목·도입·사진·답의 깊이·최신성·댓글 질문을 기록한다.",
    "공식기관·제조사·공공데이터에서 사실을 확인할 수 있는지 검토한다.",
    "브링케어가 실제로 맡을 수 있는 접수·확인·조율·보고 지점이 있는지 판정한다.",
])
h2("4.2 후보별 조사표")
add_table(["평가축", "질문", "0점", "1점", "2점"], [
    ("현재성", "지금 관심이 있는가?", "근거 없음", "계절 반복", "최근 상승/이슈"),
    ("검색 의도", "답을 찾으려는가?", "구경만", "정보 탐색", "즉시 행동"),
    ("사업 연결", "브링케어 역할이 자연스러운가?", "억지", "간접", "직접"),
    ("근거", "공식/실제 자료가 있는가?", "없음", "보조자료", "1차자료"),
    ("차별성", "상위 글의 빈틈이 있는가?", "복제", "표현 차이", "새 판단/사진"),
    ("행동", "독자가 바로 할 일이 있는가?", "없음", "저장", "점검/문의"),
], widths=[1.0, 2.45, 1.05, 1.05, 1.15], font_size=8.2)
h2("4.3 빈틈을 찾는 질문")
add_bullets([
    "상위 글이 ‘무엇’만 말하고 ‘어떤 순서’는 빼먹었는가?",
    "자가 확인과 전문가 점검의 경계가 불명확한가?",
    "한국 주거 환경·원룸·다가구 상황으로 설명되지 않았는가?",
    "사진이 예쁘기만 하고 판단 기준이 없는가?",
    "제품 추천만 있고 원인 구분이나 안전 경고가 없는가?",
    "최신 정책·제조사 안내가 반영되지 않았는가?",
])
h2("4.4 유명 이슈 연결 원칙")
add_callout("핵심", "연예인 집, 서울 아파트값, 지원금, 태풍처럼 유명한 주제는 ‘입구’가 될 수 있지만 제목에서 약속한 답을 본문에서 충분히 제공해야 합니다. 인기만 빌리고 갑자기 건물관리로 전환하면 제외합니다.", fill=YELLOW, accent=RED)

chapter(5, "주제 승인 게이트", "조회 가능성보다 사실성·사업 연결·독자 효용을 먼저 통과시킨다.")
h2("5.1 다섯 조건")
add_numbers([
    "현재 독자가 관심을 가지거나 실제로 검색하는 문제인가?",
    "브링케어 사업과 자연스럽게 연결되는가?",
    "제목의 약속을 본문에서 실제로 해결할 수 있는가?",
    "실제 사진·제공 기록·공식자료 중 하나 이상의 근거가 있는가?",
    "읽은 뒤 독자가 실행할 행동이 있는가?",
])
add_para("네 가지 이상을 충족해야 후보가 됩니다. 단, 사실성 또는 사업 연결이 실패하면 점수와 관계없이 제외합니다.")
h2("5.2 판정 결과")
add_table(["판정", "의미", "다음 행동"], [
    ("작성승인", "자료와 각도가 충분", "원고 작성"),
    ("수정후승인", "제목·검색 질문·연결을 고치면 가능", "브리프 수정 후 재검증"),
    ("검증대기", "공식근거·사진·제공기록 부족", "자료 확보"),
    ("제외", "연결이 억지이거나 안전·사실 기준 실패", "작성하지 않음"),
], widths=[1.1, 3.0, 2.6], font_size=8.8)
h2("5.3 즉시 탈락")
add_bullets(["제목과 본문 불일치", "확인되지 않은 핵심 사실", "위험한 자가 조치", "현장처럼 보이게 한 AI 이미지", "개인정보 미처리", "연예·정치·주가·사건을 인기만으로 연결"])

chapter(6, "글 유형별 설계", "콘텐츠 유형마다 기대되는 근거와 브링케어 연결 방식을 다르게 적용한다.")
add_table(["유형", "시작점", "중심 내용", "브링케어 연결", "필수 자료"], [
    ("트렌드", "계절·생활 이슈", "지금 필요한 실용 답", "여러 세대 반복 시 관련 역할", "당일 근거"),
    ("검색정보", "구체 질문", "원인·확인·행동", "현장 확인이 필요한 경계", "공식자료"),
    ("현장사례", "실제 문제", "확인·판단·처리·결과", "직접/외부/승인 구분", "실제 사진·기록"),
    ("건물주고민", "반복 부담", "비교·우선순위·흐름", "보고·승인·조율", "업무 기준"),
    ("회사소개", "관리 원칙", "한 창구·역할 경계", "한 원칙에 집중", "실제 프로세스"),
    ("수익형", "구매 전 문제", "선택 기준·사용 범위", "제품보다 관리 판단 우선", "승인 제휴링크"),
], widths=[0.8, 1.3, 1.8, 1.75, 1.2], font_size=7.8)
h2("6.1 자동화에서 제외하는 유형")
add_para("현장증명형과 건물주 고민형은 사용자가 실제 사진·작업 정보·확인 기록을 제공할 때 작성합니다. 자동화가 실적이나 고객 사례를 만들어내면 안 됩니다.")

chapter(7, "제목과 대표 이미지", "홈피드에서 즉시 이해되면서도, 검색 독자가 기대하는 답과 정확히 일치하도록 만든다.")
h2("7.1 제목 공식")
add_bullets([
    "키워드 + 실제 상황 + 숨겨진 판단: ‘에어컨 물떨어짐, 수리기사 부르기 전에 이 순서로 확인하세요’",
    "대중 이슈 + 독자 질문: ‘4인 가족 지원금, 에어컨청소에도 사용할 수 있을까?’",
    "상식 교정: ‘청소만 하면 될 줄 알았는데 공실 확인에서 더 나온 것’",
    "비교: ‘직접 처리할 때와 관리 창구가 있을 때 달라지는 일’",
    "손실 회피는 실제 근거가 있을 때만: ‘방치하면 벽지까지 번질 수 있는 물 자국’",
])
h2("7.2 제목 금지")
add_bullets(["건물주만 이해하는 내부 표현", "결론 없는 추상어", "근거 없는 숫자", "‘무조건·100%·충격’ 과장", "본문에서 해결하지 않는 유명인·정책 이름", "같은 키워드 반복"])
h2("7.3 대표 이미지 체크")
add_bullets([
    "첫 화면에서 문제 물체가 한눈에 보인다.",
    "한국 주거 환경의 실제 사진을 우선한다.",
    "텍스트는 짧고 제목을 그대로 복사하지 않는다.",
    "개인정보와 주소 단서가 없다.",
    "본문 핵심 판단과 연결된다.",
    "AI 이미지라면 실제 현장처럼 보이게 하지 않고 네이버 AI 표시를 사용한다.",
])

chapter(8, "도입부 설계", "회사 인사가 아니라 독자가 방금 겪은 장면으로 시작해 계속 읽을 이유를 만든다.")
h2("8.1 권장 4단계")
add_numbers(["실제 장면 2~4줄", "흔한 생각 또는 오해", "왜 그 생각만으로 부족한지 질문", "글에서 얻을 답 예고"])
add_callout("도입 예시", "원룸에 공실이 생기면 가장 먼저 청소를 떠올립니다. 방만 깨끗하면 바로 다음 임차인을 받을 수 있을 것 같기 때문입니다. 그런데 실제 입주 후 민원은 청소가 아니라 배수·가전·문과 창문처럼 다른 지점에서 시작되기도 합니다. 이번 글에서는 입주 전에 무엇을 먼저 구분해 확인할지 정리합니다.", fill=BEIGE)
h2("8.2 피해야 할 도입")
add_bullets(["‘안녕하세요, 브링케어입니다’로 긴 회사소개", "주제와 무관한 감상", "결론을 제목과 첫 문장에서 모두 공개", "독자가 모르는 내부 작업명", "근거 없는 공포"])

chapter(9, "본문 설득 구조", "정보 나열이 아니라 독자가 판단하고 행동할 수 있는 순서로 설명한다.")
h2("9.1 7단계 기본 구조")
add_numbers(["대상 지정", "실제 장면", "통념 교정", "근거와 원리", "판단 공식", "작은 실행", "하나의 행동"])
h2("9.2 홈피드형 구조")
add_para("대중 생활문제 → 즉시 확인할 것 → 예상과 다른 주의점 → 사진·공식근거 → 여러 세대에서 복잡해지는 지점 → 브링케어 관련 역할 → CTA")
h2("9.3 현장사례 구조")
add_para("현장 배경 → 발견 상태 → 위험·생활 영향·시급성 분류 → 건물주 보고·승인 → 직접 관리/외부 작업 → 완료 사진 → 다음 관리 기준")
h2("9.4 반복을 줄이는 법")
add_bullets([
    "같은 뜻의 ‘확인하고 전달하고 조율합니다’를 여러 소제목에서 반복하지 않는다.",
    "서비스 설명은 글 주제와 닿는 한 문단으로 제한한다.",
    "한 문단이 없어도 결론이 같다면 삭제한다.",
    "목록 뒤에 같은 내용을 문장으로 다시 쓰지 않는다.",
])

chapter(10, "네이버 모바일 편집 규칙", "작성 원고가 실제 네이버 화면에서 읽히도록 정렬·여백·강조·컴포넌트를 통일한다.")
h2("10.1 정렬")
add_callout("고정 규칙", "제목을 제외한 본문은 가운데 정렬을 기본으로 합니다. 긴 표·복잡한 단계처럼 가운데 정렬이 오히려 읽기 어려운 특별한 경우만 제한적으로 예외를 둡니다.", fill=SKY)
h2("10.2 문장과 빈 문단")
add_bullets([
    "한 문장 또는 하나의 의미 묶음이 끝나면 네이버 편집기에서 빈 문단 하나를 넣는다.",
    "같은 설명을 잇는 짧은 1~3문장은 붙일 수 있다.",
    "대상·상황·판단·행동이 바뀌는 지점에는 반드시 한 줄을 더 띄운다.",
    "문장마다 기계적으로 띄우지 말고 의미 블록 사이에 여백을 만든다.",
    "빈 글머리표나 점만 남는 문단을 만들지 않는다.",
])
h2("10.3 실제 편집 예시")
add_callout("좋은 호흡", "여러 세대를 관리한다면 기록 방식도 같아야 합니다.\n\n원룸이나 다가구 건물은 공실이 생길 때마다 건물주가 직접 방문하기 어렵습니다.\n\n브링케어는 현장을 확인하고 사진으로 상태를 기록한 뒤, 청소로 해결할 일과 전문업체 확인이 필요한 일을 구분합니다.", fill=LIGHT)
h2("10.4 강조와 색깔 밑줄")
add_bullets([
    "색깔 밑줄처럼 보이는 효과는 네이버 편집기의 글자 배경색을 사용한다.",
    "연두: 변화·해결 포인트, 베이지: 판단 기준, 노랑: 주의·핵심 행동.",
    "한 문단 전체가 아니라 핵심 구절 1개만 칠한다.",
    "한 화면에 여러 색을 섞지 않는다.",
    "굵게·밑줄·배경색을 같은 구절에 과도하게 중첩하지 않는다.",
])
h2("10.5 인용구·구분선·이모티콘")
add_bullets([
    "인용구: 글의 핵심 판단 또는 전환 문장에 사용.",
    "구분선: 도입→해결, 정보→서비스 연결처럼 큰 블록 전환에 사용.",
    "문자로 만든 ‘────’는 사용하지 않고 네이버 실제 구분선 컴포넌트를 쓴다.",
    "이모티콘: 소제목이나 감정 전환에 1개 정도. 같은 이모티콘 반복 금지.",
])

chapter(11, "사진 수집·저작권·개인정보", "실제 사진의 신뢰를 살리면서 무단 사용·개인정보·기만 위험을 제거한다.")
h2("11.1 사진 우선순위")
add_numbers(["브링케어가 직접 촬영한 실제 현장 사진", "사용자가 제공한 사진", "공식기관·제조사의 재사용 가능한 자료", "대한민국 공공누리·공공저작물", "명확한 라이선스의 실사진", "설명도 또는 AI 이미지(최후 수단)"])
h2("11.2 사진 선택 원칙")
add_bullets(["사진 수는 고정하지 않는다.", "각 사진에 ‘도입·원인·확인·결과·CTA’ 중 역할이 있어야 한다.", "비슷한 사진 여러 장보다 정보가 다른 사진을 고른다.", "한국 주거 환경 사진을 우선한다.", "대표 사진 한 장에 모든 정보를 몰지 않는다."])
h2("11.3 개인정보 가림")
add_table(["가림 대상", "예시", "처리"], [
    ("위치", "상세 주소, 동·호수, 우편물", "크롭 또는 블러"),
    ("출입", "비밀번호, 공동현관 정보, 열쇠 번호", "완전 삭제"),
    ("사람", "얼굴, 명찰, 개인 연락처", "동의 없으면 블러"),
    ("차량", "번호판, 주차증", "블러"),
    ("문서", "계약서, 영수증, 계좌·금액", "공개 허용 범위만 남김"),
], widths=[1.1, 2.5, 3.1], font_size=8.6)
h2("11.4 출처 표기 방식")
add_para("사진의 저작자·라이선스·원본 URL은 내부 사진 권리 원장에 보관합니다. 공개 본문에는 독자 흐름을 방해하는 긴 출처 덤프를 넣지 않습니다. 단, 라이선스가 공개 표시를 요구하면 사진 가까운 캡션에 필요한 범위로 표기합니다.")

chapter(12, "공식 자료 활용과 출처 처리", "독자가 외부 링크를 눌러야만 답을 알 수 있는 글을 만들지 않으면서도 사실을 정확히 검증한다.")
h2("12.1 공식 자료 사용법")
add_numbers(["실제 공식 페이지를 연다.", "확인 날짜와 문서 날짜를 기록한다.", "핵심 안내를 이해한 뒤 브링케어 독자 상황에 맞게 재구성한다.", "원문 문장을 길게 복사하지 않는다.", "제품·모델별 차이가 있으면 범위를 명시한다.", "독자가 이 글 안에서 필요한 답을 얻도록 충분히 설명한다."])
h2("12.2 출처의 공개 위치")
add_bullets([
    "정책 신청·가격·안전처럼 독자가 직접 확인해야 하는 링크는 관련 문단에 자연스럽게 연결한다.",
    "일반 검증용 출처는 내부 원장에 보관한다.",
    "본문 맨 아래 ‘공식 참고자료 / 확인 기준일 / URL 목록’을 자동으로 붙이지 않는다.",
    "사진 설명이 필요하면 사진 가까이에 짧게 출처를 표시한다.",
])
h2("12.3 최신성 필수 분야")
add_bullets(["정책·법률·지원사업", "가격·요금·제품 사양", "날씨·태풍·재난", "안전·화재·전기·가스", "부동산 통계·청약·거래", "제조사 관리 방법"])

chapter(13, "브링케어 연결 문단과 CTA", "정보 글을 회사소개로 망치지 않고 독자의 다음 행동만 자연스럽게 연결한다.")
h2("13.1 연결 원칙")
add_bullets([
    "정보 답변을 먼저 완결한다.",
    "주제와 관련된 브링케어 역할만 1개 의미 블록으로 설명한다.",
    "모든 글에 동일한 긴 회사소개를 붙이지 않는다.",
    "브링케어가 직접 하지 않는 전문 작업을 명확히 구분한다.",
    "한 글의 CTA는 카카오채널 또는 전화상담 중 하나만 선택한다.",
])
h2("13.2 연결 공식")
add_callout("서비스 연결", "여러 세대에서 같은 문제가 반복될 때는 누가 호실·증상·사진을 모으고, 현장 확인이 필요한지 구분하며, 건물주 승인과 업체 일정을 조율할지가 중요합니다. 브링케어는 이 관리 과정을 한 창구로 연결합니다.", fill=MINT)
h2("13.3 CTA 공식")
add_para("‘건물 위치 / 층수·세대수 / 문제가 발생한 공간 / 보이는 증상 / 사진’을 카카오채널 BRING Care로 보내주세요. 확인 가능한 범위에서 청소·현장 확인·전문업체 점검 중 다음 단계를 안내합니다.")
h2("13.4 최종 배너")
add_para("승인된 상담 배너는 글 맨 마지막에 한 번만 삽입합니다. 본문에 전화·유튜브·인스타그램을 다시 나열하지 않습니다. 배너는 강한 CTA이자 연락처 서명 모듈입니다.")

chapter(14, "현장 작업 글 작성법", "사진과 작업 기록을 받은 뒤 현장을 과장 없이 증거형 콘텐츠로 전환한다.")
h2("14.1 사용자가 제공해야 할 최소 정보")
add_bullets(["작업 날짜와 대략적 위치 범위", "무엇을 발견했는지", "누가 작업했는지", "건물주 승인 내용", "전·중·후 사진", "완료 확인 방식", "비용·시간 공개 가능 여부", "가려야 할 정보"])
h2("14.2 사진 판독 원칙")
add_bullets(["모든 로컬 사진을 직접 열어 본다.", "사진에서 보이는 사실과 사용자가 설명한 사실을 구분한다.", "사진만으로 원인·안전·완료를 단정하지 않는다.", "전후 사진의 촬영 방향이 다르면 비교 한계를 적는다."])
h2("14.3 현장 글 구성")
add_numbers(["독자가 공감할 일반 문제", "이번 현장의 공개 가능한 배경", "사진에서 확인된 상태", "왜 우선순위를 정했는지", "브링케어 직접 역할", "외부업체 작업과 건물주 승인", "완료 사진과 남은 항목", "독자가 자신의 건물에서 확인할 체크리스트", "상담 CTA"])
h2("14.4 금지")
add_bullets(["없는 고객 반응 만들기", "확인하지 않은 원인 단정", "작업자를 브링케어 직원으로 오인", "비용 공개 동의 없이 금액 노출", "상세 주소·호실 공개", "결과를 과장한 전후 비교"])

chapter(15, "검색 해결형 글 작성법", "한 검색 질문에 이 글 하나로 충분한 답을 제공한다.")
h2("15.1 답의 구조")
add_numbers(["증상 또는 질문 정의", "흔한 원인 범주", "안전한 자가 확인", "하지 말아야 할 조치", "청소·수리·전문점검 경계", "반복·다세대 상황", "브링케어 관련 역할", "CTA"])
h2("15.2 깊이 기준")
add_bullets(["외부 링크를 눌러야 핵심 답을 알 수 있게 쓰지 않는다.", "제품마다 다른 부분은 모델 확인법을 알려준다.", "자가 조치가 위험한 분야는 전원 차단·전문가 문의 등 안전 기준을 우선한다.", "목록만 나열하지 않고 왜 그 순서인지 설명한다."])

chapter(16, "대중 유입형·홈피드형 글 작성법", "잘 알려진 이슈를 입구로 사용하되 브링케어와의 연결을 억지로 만들지 않는다.")
h2("16.1 홈피드에서 중요한 것")
add_bullets(["첫눈에 이해되는 제목", "대표 이미지의 즉시성", "도입 5~8줄 안의 문제 약속", "예상과 다른 판단 포인트", "저장 가능한 체크리스트", "실제 사진 또는 공식 근거", "짧고 자연스러운 사업 연결"])
h2("16.2 공개되지 않은 알고리즘에 대한 태도")
add_para("글자 수, 사진 수, 발행 시간, 키워드 반복 횟수, 해시태그 수를 순위 공식처럼 단정하지 않습니다. 실제 성과를 72시간·7일·14일·30일 단위로 기록해 우리 계정의 패턴을 학습합니다.")
h2("16.3 홈피드형 제목 검수")
add_bullets(["한 번에 이해되는가?", "궁금증이 남는가?", "과장 없이 클릭 이유가 있는가?", "대표 이미지와 같은 약속인가?", "본문에서 답을 얻는가?"])

chapter(17, "부동산·연예인 공간·시사 인접 주제", "대중성이 큰 분야를 안전하고 자연스럽게 주거관리 관점으로 다룬다.")
h2("17.1 부동산")
add_bullets(["한국부동산원·국토교통부·서울시·청약홈 등 1차 자료 우선", "집값 상승·하락을 투자 조언으로 단정하지 않음", "거래 통계와 개별 단지 사례를 구분", "입주·공실·관리비·주거 상태처럼 실제 관리 행동으로 연결"])
h2("17.2 연예인 집·인테리어")
add_bullets(["공식 방송·공식 채널에서 공개된 사실만 사용", "나무위키는 관련어 탐색용일 뿐 근거로 사용하지 않음", "주소·재산·집값·사생활 추정 금지", "무단 캡처·사진 사용 금지", "공간 구성에서 일반 독자가 적용할 수 있는 관리·수납·환기 관점만 추출"])
h2("17.3 시사·지원정책")
add_para("지원금·정책이 실제로 유효한 기간인지 먼저 확인합니다. 과거에 끝난 제도를 현재 신청 가능한 것처럼 제목에 쓰지 않습니다. 검색량이 남아 있어도 현재 독자가 얻을 수 있는 답이 없다면 과거 제도 정리 또는 대체 지원 확인법으로 각도를 바꿉니다.")

chapter(18, "쿠팡 파트너스 수익형", "수익을 위해 신뢰를 희생하지 않고, 승인된 링크와 실제 구매 판단 정보만 사용한다.")
h2("18.1 적용 조건")
add_bullets(["사용자가 제공하거나 승인한 제휴 링크만 사용", "주제와 직접 관련된 상품만 배치", "경제적 이해관계 고지를 명확히 표시", "가격·후기·효과·재고를 임의로 생성하지 않음", "상품보다 문제 해결 기준을 먼저 제공"])
h2("18.2 승인 링크 원장")
add_table(["상품", "일반 링크", "배너 링크", "상태"], [
    ("욕실 클리너", "https://link.coupang.com/a/gcPwTAYCaG", "https://link.coupang.com/a/gcPx5COS8O", "사용자 승인"),
    ("세탁조 클리너", "https://link.coupang.com/a/gc3UXOQedM", "https://link.coupang.com/a/gc3Vqj1BbU", "사용자 승인"),
], widths=[1.2, 2.25, 2.25, 1.0], font_size=8)
h2("18.3 수익형 글 구성")
add_numbers(["문제 상황", "제품을 쓰기 전에 확인할 원인", "제품이 맞는 범위와 맞지 않는 범위", "사용설명서 우선", "승인 제휴 고지와 링크", "사용 후 확인", "여러 세대 관리 시 기록·조율", "CTA"])
h2("18.4 API 없이 운영할 때")
add_para("쿠팡 파트너스 API 자격이 없으면 승인 링크를 수동 원장으로 관리합니다. 링크 유효성은 발행 전 열어 확인하고, 상품명·가격·이미지를 자동 생성하지 않습니다. API 미사용은 자동화 실패가 아니라 링크 승인 단계를 사람이 맡는 운영 방식입니다.")

chapter(19, "금지 표현·안전·사실성", "조회수와 상담을 위해 사실을 왜곡하거나 위험한 행동을 권하지 않는다.")
h2("19.1 근거 없이 보장하면 안 되는 것")
add_bullets(["공실 감소", "비용 절감", "응답 시간", "수리 성공", "검색·홈피드 순위", "특정 제품의 치료·살균·완전 제거 효과"])
h2("19.2 위험 분야")
add_table(["분야", "자가 안내 가능 범위", "전문가로 넘길 기준"], [
    ("전기", "외관·차단기 상태 확인, 전원 분리", "타는 냄새·스파크·반복 차단"),
    ("가스", "사용 중단·환기·공식 신고 안내", "냄새·누출 의심 즉시"),
    ("누수", "물이 보이는 위치·시간·사진 기록", "지속 누수·전기부 인접·천장 팽창"),
    ("곰팡이", "습기 원인·환기·보호장비", "광범위·반복·누수 동반"),
    ("가전", "모델명·증상·외부 상태 확인", "분해·냉매·내부 배선"),
], widths=[0.9, 3.4, 2.4], font_size=8.2)
h2("19.3 숫자 사용")
add_para("제목의 횟수·금액·기간·비율은 공식 근거 또는 실제 제공 기록이 있을 때만 씁니다. 보기 좋게 만들기 위해 ‘5가지’, ‘3단계’를 임의로 고정하지 않습니다.")

chapter(20, "발행 전 검증", "작성자의 감각이 아니라 동일한 게이트와 체크리스트로 품질을 확인한다.")
h2("20.1 브리프 검증")
add_bullets(["요청 모드", "배포 목표", "주역할", "글 유형", "주제·주키워드", "독자·장면·불안·약속", "검증 사실", "업무 범위", "사진", "CTA", "공개 위치 범위", "비용·시간 공개 여부"])
h2("20.2 원고 검증")
add_bullets(["제목과 본문 일치", "근거 없는 주장 없음", "공식 URL 유효", "개인정보 없음", "브랜드 연락처 정확", "CTA 하나", "배너 한 번", "미확인 항목 없음", "AI·실제 사진 구분", "키워드 부자연스러운 반복 없음"])
h2("20.3 사람의 최종 판단")
add_para("자동 검증 통과는 발행 품질의 최소선입니다. 실제 네이버 화면의 정렬·여백·사진 순서·강조·구분선·링크·카테고리를 사람이 다시 확인합니다.")

chapter(21, "네이버 편집기 발행 절차", "원고를 붙여 넣는 데서 끝나지 않고 네이버 전용 컴포넌트와 설정을 완성한다.")
add_numbers([
    "로그인 상태와 작성 계정을 확인한다.",
    "기존 공개 글과 backlog에서 중복을 확인한다.",
    "제목을 입력하고 본문 전체를 가운데 정렬한다.",
    "의미 묶음 사이에 실제 빈 문단을 넣는다.",
    "핵심 구절에만 연두·베이지·노랑 배경색을 넣는다.",
    "문자선이 아니라 실제 인용구·구분선을 삽입한다.",
    "실제 사진을 관련 문단 가까이에 배치하고 캡션을 넣는다.",
    "승인된 상담 배너를 맨 마지막에 한 번 삽입한다.",
    "태그는 본문이 아니라 네이버 태그 입력란에 등록한다.",
    "카테고리·전체공개·검색 허용을 확인한다.",
    "모바일 미리보기에서 문장 깨짐과 여백을 확인한다.",
    "발행 후 공개 페이지를 다시 연다.",
])
h2("21.1 하단에 넣지 않는 블록")
add_bullets(["태그 제목과 해시태그 나열", "공식 참고자료 URL 덤프", "확인 기준일 단독 블록", "사진 출처·이용 조건 장문", "브링케어 연락처 반복 목록", "네이버가 자동 생성한 출처·작성자 문구의 수동 복제"])

chapter(22, "발행 후 공개 검수", "편집 화면이 아니라 실제 공개 페이지를 기준으로 완료를 판정한다.")
add_table(["검수 항목", "확인 방법", "실패 시 조치"], [
    ("제목", "잘림·오탈자·약속 일치", "즉시 수정"),
    ("정렬", "제목 제외 본문 가운데 정렬", "전체 선택 후 재정렬"),
    ("여백", "의미 블록 사이 빈 문단 표시", "문단별 보완"),
    ("사진", "순서·중복·개인정보·캡션", "교체/가림"),
    ("강조", "색 배경이 핵심 구절에만 적용", "과다 강조 제거"),
    ("구분선", "실제 컴포넌트인지", "문자선 교체"),
    ("CTA/배너", "CTA 하나, 배너 마지막 1회", "중복 삭제"),
    ("설정", "카테고리·전체공개·검색허용·태그", "설정 수정"),
], widths=[1.0, 3.5, 2.2], font_size=8.4)

chapter(23, "성과 측정과 개선", "감이 아니라 글별 데이터를 축적해 브링케어 계정에서 실제로 반응하는 패턴을 찾는다.")
h2("23.1 측정 시점")
add_table(["시점", "확인", "판단"], [
    ("72시간", "초기 조회·유입 경로·표지 반응", "제목/대표이미지 문제"),
    ("7일", "검색어·체류·공감·댓글", "답의 깊이와 검색 의도"),
    ("14일", "검색 유입 지속·관련글 이동", "롱테일 가능성"),
    ("30일", "누적 조회·상담·재활용 가치", "확장/수정/중단"),
], widths=[1.0, 3.1, 2.6], font_size=8.7)
h2("23.2 글별 기록 항목")
add_bullets(["제목·주키워드·글 유형·역할", "발행일·카테고리", "대표 이미지 유형", "조회·유입 경로", "검색어", "평균 체류 또는 읽기 반응", "공감·댓글·저장", "상담 클릭·문의", "수정 내용", "다음 실험"])
h2("23.3 개선 규칙")
add_bullets(["노출이 적다: 주제 수요·검색어·카테고리 재검토", "노출은 있으나 클릭이 적다: 제목·대표이미지 수정", "클릭은 있으나 이탈이 빠르다: 도입 약속·본문 깊이 수정", "읽지만 상담이 없다: 서비스 연결과 CTA 적합성 수정", "상담은 있으나 부적합하다: 대상·범위 문구 명확화"])

chapter(24, "자동화 운영", "자동화가 품질을 희생하지 않도록 상태 원장과 중단 조건을 명확히 한다.")
h2("24.1 상태 단계")
add_para("후보 → 조사완료 → 본문완료 → 이미지보완 → 검증완료 → 승인대기 → 발행완료")
h2("24.2 backlog 원칙")
add_bullets(["미완료 작업을 버리지 않는다.", "매 회차 기존 공개 글과 backlog를 먼저 읽는다.", "이미지가 부족해도 검증된 본문을 폐기하지 않고 이미지보완으로 남긴다.", "제목·핵심 검색어·문제·결론·연결 문장 중복을 확인한다.", "출처·확인일·사진 권리를 내부 원장에 기록한다."])
h2("24.3 자동 발행 중단 조건")
add_bullets(["로그인 만료", "CAPTCHA", "계정·정책 경고", "공식 근거 부족", "사진 권리 불명", "개인정보 미처리", "사업 연결 약함", "기존 글과 실질 중복", "검증 오류", "현장 사실 부족"])
h2("24.4 자동화에서 사람이 맡는 일")
add_bullets(["현장 사진과 작업 기록 제공", "가격·비용·고객 발언 공개 승인", "새 쿠팡 링크 승인", "권리 불명 사진 사용 판단", "정책 경고 대응", "브랜드 방향의 큰 변경"])

chapter(25, "실패 사례와 교정법", "지금까지 반복된 문제를 재발 방지 규칙으로 전환한다.")
failures = [
    ("글이 너무 짧음", "목록만 있고 이유·순서·경계가 없음", "독자가 이 글 안에서 답을 얻도록 원리·판단·행동을 보강"),
    ("AI스럽고 딱딱함", "정리체·같은 어미·추상어 반복", "실제 장면→질문→판단→행동으로 재작성, 짧은 구어체 문단"),
    ("여백 없음", "문장이 한 덩어리", "의미 묶음이 끝날 때 실제 빈 문단 삽입"),
    ("가운데 정렬 누락", "일부만 중앙, 나머지 좌측", "제목 제외 본문 전체를 공개 페이지에서 확인"),
    ("출처 덤프", "맨 아래 URL·라이선스·확인일 장문", "내부 원장 보관, 필요한 링크만 관련 문단/캡션"),
    ("사진이 적거나 무관", "대표 이미지 외 설명 사진 없음", "도입·원인·판단·결과 가까이에 역할별 배치"),
    ("AI 사진 오인", "현장처럼 보이는 생성 이미지", "실제 한국 사진 우선, AI 표시와 역할 구분"),
    ("서비스 설명 반복", "모든 글 끝에 같은 장문", "주제 관련 역할만 한 의미 블록"),
    ("핫이슈 억지 연결", "유명 제목 후 갑자기 건물관리", "사업 연결 게이트 실패 시 제외"),
    ("외부 링크 의존", "답은 없고 제조사 사이트로 이동 요구", "핵심 내용을 충분히 재구성하고 링크는 검증·보조"),
    ("글머리표 잔재", "빈 점 여러 개", "공개 페이지에서 빈 리스트·문단 제거"),
    ("문자 구분선", "──── 사용", "네이버 실제 구분선 컴포넌트로 교체"),
]
add_table(["실패", "원인", "교정"], failures, widths=[1.25, 2.2, 3.35], font_size=7.8)

chapter(26, "표준 입력 양식", "사용자는 일한 내용과 사진만 제공해도 작성자가 누락 없이 검증 가능한 브리프를 만들 수 있게 한다.")
h2("26.1 현장 작업 제공 양식")
for label in ["작업 날짜", "공개 가능한 위치 범위", "건물 유형", "발견한 문제", "건물주 요청", "브링케어가 직접 한 일", "외부업체가 한 일", "건물주가 승인한 일", "작업 전 상태", "작업 중 과정", "작업 후 상태", "남은 문제", "비용 공개 가능 여부", "시간 공개 가능 여부", "가릴 정보", "원하는 CTA"]:
    add_para(f"{label}: ______________________________________________", after=4)
h2("26.2 사진 목록표")
add_table(["번호", "파일명", "촬영 시점", "보이는 사실", "본문 역할", "가림", "캡션"], [[str(i), "", "", "", "", "", ""] for i in range(1, 9)], widths=[0.45, 1.2, 0.8, 1.5, 1.0, 0.7, 1.2], font_size=7.4)
h2("26.3 키워드 브리프")
for label in ["요청 모드", "배포 목표", "콘텐츠 주역할", "글 유형", "주제", "주키워드", "보조키워드", "목표 독자", "독자 장면", "독자 불안", "약속할 답", "CTA", "공식 출처", "확인 날짜", "기존 글 중복"]:
    add_para(f"{label}: ______________________________________________", after=4)

chapter(27, "표준 출력 패키지", "초안뿐 아니라 제목·사진·태그·검수·성과표까지 발행에 필요한 전체 묶음을 동일하게 산출한다.")
add_numbers([
    "승인 판정·이유·부족자료",
    "독자·상황·불안·약속·행동 한 줄 요약",
    "내부 제목 12개 중 상위 3개와 추천 1개",
    "네이버 게시용 최종 본문",
    "인용구·구분선·사진 위치",
    "사진 순서·캡션·실제/AI 표시·가림",
    "추가 필요 사진",
    "공식 출처·확인 기준일(내부 원장)",
    "주제 맞춤 브링케어 연결 문단",
    "선택 CTA와 사용자가 보낼 정보",
    "썸네일 문구 3개",
    "검색어·네이버 태그·선택적 클립 대본",
    "사실성·안전·개인정보·반복 검수",
    "확인 필요 항목",
    "72시간·7일·14일·30일 성과 회고표",
])

chapter(28, "체크리스트·기록지·개정 이력", "매번 사용할 수 있는 최종 체크리스트와 이 매뉴얼의 변경 내역을 유지한다.")
h2("28.1 발행 전 체크리스트")
checks = [
"주제가 현재 관심 또는 실제 검색 문제다.", "사업 연결이 자연스럽다.", "제목과 본문 약속이 일치한다.", "핵심 사실을 공식 또는 제공 기록으로 확인했다.", "현장 사례를 지어내지 않았다.", "브링케어 직접·외부업체·건물주 결정을 구분했다.", "상세 주소·호실·얼굴·차량번호·연락처를 가렸다.", "제목 제외 본문이 가운데 정렬이다.", "의미 블록 사이에 실제 빈 문단이 있다.", "인용구와 구분선은 실제 컴포넌트다.", "색 배경 강조가 핵심 구절에만 있다.", "사진은 실제 한국 사진을 우선했고 역할이 있다.", "사진 권리와 출처를 내부 원장에 기록했다.", "출처·태그·연락처 덤프가 본문 하단에 없다.", "태그는 네이버 태그란에 입력한다.", "CTA는 하나다.", "상담 배너는 마지막에 한 번만 있다.", "카테고리·전체공개·검색허용을 확인했다.", "모바일 미리보기를 확인했다.", "발행 후 공개 페이지를 다시 열어 확인했다."
]
for c in checks:
    add_para(f"☐ {c}", after=3)
h2("28.2 발행 후 회고 기록지")
add_table(["항목", "72시간", "7일", "14일", "30일"], [
    ("조회수", "", "", "", ""), ("주요 유입", "", "", "", ""), ("검색어", "", "", "", ""), ("공감/댓글", "", "", "", ""), ("상담", "", "", "", ""), ("수정", "", "", "", ""), ("다음 실험", "", "", "", ""),
], widths=[1.3, 1.35, 1.35, 1.35, 1.35], font_size=8)
h2("28.3 개정 이력")
add_table(["버전", "날짜", "변경 내용", "변경 이유", "영향 범위"], [
    ("1.0", "2026-08-17", "전체 운영 체계 최초 통합: 전략·조사·작성·사진·편집·검증·발행·성과·자동화", "대화와 개별 문서에 흩어진 규칙을 단일 기준서로 통합", "전체"),
    ("1.1", "2026-08-17", "장애 알림, AI 이미지 대체, 성과 학습 루프 추가", "무인 자동화의 실패 복구와 지속 개선", "자동화·이미지·성과·매뉴얼"),
    ("", "", "", "", ""), ("", "", "", "", ""), ("", "", "", "", ""),
], widths=[0.7, 1.05, 2.45, 1.7, 0.9], font_size=7.8)
h2("28.4 새 규칙을 추가하는 방법")
add_numbers(["새 합의가 기존 어느 장에 속하는지 찾는다.", "기존 문구와 충돌하는지 확인한다.", "규칙·예시·검수 항목을 함께 수정한다.", "개정 이력에 날짜·사유·영향 범위를 기록한다.", "자동화나 검증기에 영향이 있으면 함께 업데이트한다.", "다음 실제 글에서 적용하고 공개 페이지에서 검증한다."])

chapter(29, "학습형 자동화 운영", "장애를 안전하게 멈추고, 이미지 대체와 실제 성과를 다음 실행에 반영한다.")
h2("29.1 로그인·CAPTCHA·편집기 장애 알림")
add_bullets(["로그인 만료·CAPTCHA·편집기 구조 변경·정책 경고·공개 검수 실패가 확인되면 즉시 클릭과 추가 발행을 멈춘다.", "알림에는 발생 시각, 글 제목, 중단 단계, 사용자가 해야 할 조치, 해결 후 재개 지점을 함께 적는다.", "원고와 이미지는 폐기하지 않고 보존하며, 상태 변화가 없는 동일 장애는 24시간 안에 반복 알림하지 않는다."])
h2("29.2 조건부 AI 이미지 전환")
add_bullets(["대한민국 실사진과 재사용 권리가 확인된 자료를 먼저 찾는다.", "검색정보·트렌드 글은 실사진 확보 실패 이유를 기록하고 독자 오인이 없을 때만 설명용 AI 이미지로 전환한다.", "AI 이미지는 짧게 AI 활용 이미지임을 표시한다.", "현장사례·전후 비교·고객 작업 기록에는 AI 이미지를 사용하지 않는다."])
h2("29.3 72시간·7일·14일·30일 성과 수집")
add_table(["시점", "기록", "활용"], [("72시간", "조회·유입·초기 반응", "제목·대표 이미지 진단"), ("7일", "검색어·체류·반응·상담", "검색 의도와 답의 깊이 진단"), ("14일", "검색 지속·관련 이동·상담", "롱테일과 확장 판단"), ("30일", "누적 조회·상담·제휴 행동", "유지·수정·쿨다운 판단")], widths=[1.0, 3.0, 2.7], font_size=8.5)
add_para("확인할 수 없는 수치는 0으로 추정하지 않고 NA로 기록한다. 20개 미만의 비교 집단은 잠정 기준으로만 쓰고, 20개 이상부터 중앙값을 기준으로 사용한다.")
h2("29.4 성과 기반 다음 주제 점수")
add_table(["항목", "배점"], [("현재 관심도", "20"), ("검색·구매·상담 의도", "20"), ("브링케어 사업 연관성", "20"), ("근거·이미지 준비도", "15"), ("기존 글과 차별성", "15"), ("과거 유사 콘텐츠 성과", "10")], widths=[5.4, 1.3], font_size=8.8)
add_para("과거 성과는 최대 10점만 반영한다. 사실성·사업 연관성·제목과 본문 일치·개인정보와 저작권·독자 행동 안전성 중 하나라도 실패하면 총점과 관계없이 제외한다.")
h2("29.5 주제 쿨다운과 실험 원장")
add_bullets(["같은 주제가 비교 기준의 절반 이하인 TOPIC_WEAK로 세 번 반복되면 60일 쿨다운한다.", "조회가 적어도 상담이 발생한 글은 전환 성공으로 분리해 주제 실패로 단정하지 않는다.", "실험은 제목·도입·대표 이미지·CTA 중 한 변수만 바꾸고 성공 지표와 안전 지표를 함께 기록한다."])
h2("29.6 매뉴얼 개정 후보")
add_para("동일한 교정이 서로 다른 글에서 세 번 이상 반복될 때만 매뉴얼 개정 후보를 만든다. 자동화는 후보와 근거를 기록할 수 있지만 Word 매뉴얼·스킬·자동화 지시문을 스스로 바꾸지 않는다. 사용자가 승인한 뒤 세 자산에 동시에 반영한다.")

doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_p(p, before=110, after=14)
r = p.add_run("끝이 아니라, 다음 개정의 시작")
set_font(r, 20, True, NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_p(p, after=22)
r = p.add_run("새로운 시행착오가 생기면 대화에만 남기지 말고\n이 매뉴얼의 규칙·예시·체크리스트에 반영합니다.")
set_font(r, 11, False, GRAY)
add_callout("브링케어 콘텐츠의 최종 기준", "독자가 필요한 답을 충분히 얻고, 실제 근거를 확인할 수 있으며, 브링케어가 맡는 역할과 맡지 않는 역할이 분명하고, 모바일 화면에서 편하게 읽히는 글을 만든다.", fill=MINT, accent=NAVY)

# Global cleanup and metadata
doc.core_properties.title = "브링케어 네이버 블로그 마스터 운영 매뉴얼"
doc.core_properties.subject = "시장조사, 콘텐츠 기획, 작성, 사진, 편집, 검증, 발행, 성과 개선"
doc.core_properties.author = "BRING Care"
doc.core_properties.keywords = "브링케어, 네이버 블로그, 운영 매뉴얼, 콘텐츠 마케팅, 건물관리"

doc.save(OUT)
print(OUT)
